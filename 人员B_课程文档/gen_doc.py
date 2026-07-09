# -*- coding: utf-8 -*-
"""生成人员B（页面基础渲染/表单/父子·跨级组件通信）课程说明 Word 文档。
所有实现均映射到本项目真实文件/组件。"""
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUT = "人员B_知识点与实现说明.docx"

doc = Document()

# 中文字体（解决 docx 中文显示）
style = doc.styles["Normal"]
style.font.name = "Microsoft YaHei"
style.font.size = Pt(11)
style.element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")


def h(text, level):
    p = doc.add_heading(text, level=level)
    for r in p.runs:
        r.font.name = "Microsoft YaHei"
        r._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    return p


def para(text, bold=False, italic=False, size=11, color=None):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold
    r.italic = italic
    r.font.size = Pt(size)
    r.font.name = "Microsoft YaHei"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    if color:
        r.font.color.rgb = RGBColor(*color)
    return p


def bullet(text, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    if bold_prefix:
        rb = p.add_run(bold_prefix)
        rb.bold = True
        rb.font.name = "Microsoft YaHei"
        rb._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    r = p.add_run(text)
    r.font.name = "Microsoft YaHei"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    return p


# ============ 封面 ============
title = doc.add_heading("电子微商城项目 — 人员B 负责模块说明", level=0)
for r in title.runs:
    r.font.name = "Microsoft YaHei"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
sub = para("页面基础渲染 · 表单开发 · 父子 / 跨级组件通信", italic=True, size=12, color=(0x88, 0x88, 0x88))
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER

# ============ 一、知识点对照表 ============
h("一、知识点对照（说明 + 对应实现）", 1)
para("以下 8 个考核知识点，均在本项目中有真实落地代码。", italic=True, color=(0x66, 0x66, 0x66))

rows = [
    ("响应式数据\n(ref / reactive / computed)",
     "用 ref 声明基本类型/引用状态，reactive 声明响应式对象，computed 派生状态；数据变化自动驱动视图更新。",
     "ProductDetail.vue：const product = computed(...)、currentImage = ref(0)、quantity = ref(1)；\nOrderList.vue：orders = ref([])、activeTab = ref('')；\n各 Pinia store（cart/product/auth）用 ref/computed 管理状态。"),
    ("内容/属性绑定\n({{}} / v-text / v-bind)",
     "文本插值 {{}} 安全渲染内容（等价 v-text，故本项目不用 v-html 以免 XSS）；v-bind(:) 动态绑定 class/style/属性。",
     "ProductDetail.vue：:src=\"product.images?.[currentImage] || product.image\"、:class=\"{active: idx===currentImage}\"；\nOrderList.vue：:class=\"{active: activeTab===tab.key}\"；\nAppHeader.vue：v-bind(\"siteConfig?.themeColor\") 动态主题色实时生效。"),
    ("事件绑定\n(v-on / 修饰符 / 事件对象)",
     "@click 等绑定事件；.enter/.prevent/.stop 为事件修饰符；事件对象 $event 可直接操作 DOM。",
     "Cart.vue：@change=\"cartStore.toggleAll(($event.target).checked)\"（读复选框状态）；\nSearchBar.vue：@keyup.enter=\"handleSubmit\"（回车提交修饰符）；\nProductDetail.vue：@click=\"quantity--\" 内联事件。"),
    ("双向绑定\n(v-model 表单)",
     "表单元素 v-model 与数据双向同步；组件上 v-model = :modelValue + @update:modelValue。",
     "SearchBar.vue：v-model=\"keyword\" 对外暴露 v-model（props modelValue + emit update:modelValue）；\nLogin.vue 登录表单、AddressPanel.vue 地址新增/编辑、ProfilePanel.vue 资料编辑、ProductDetail.vue 评价表单 v-model=\"reviewForm.content\"。"),
    ("条件 & 列表渲染\n(v-if/v-show/v-for + EP 表格)",
     "v-for 渲染列表；v-if/v-show 控制显隐；Element Plus el-table 渲染表格数据。",
     "OrderList.vue：<el-table :data=\"order.items\"> + v-for=\"order in orders\" + v-for=\"tab in tabs\"；\nCart.vue：v-for=\"item in cartStore.items\"；\nProductDetail.vue：v-for 轮播图/评价列表；Home.vue 商品列表。"),
    ("父子组件传参\n(props 校验 + 自定义事件)",
     "父传子用 props 并做类型/默认值校验；子传父用 defineEmits 触发自定义事件。",
     "StarRating.vue：defineProps({rating:{type:Number,default:0}, max:{type:Number,default:5}})；\nCartBadge.vue：defineProps({count:{type:Number,default:0}})；\nSearchBar.vue：defineEmits(['update:modelValue','search'])，父组件 @search=\"onSearch\" 接收。"),
    ("provide / inject\n跨级传值",
     "祖先 provide 数据，任意后代 inject 获取，避免逐层 props 透传。",
     "App.vue：provide('siteConfig', siteConfig) 与 provide('updateSiteConfig', updateSiteConfig)；\nAppHeader.vue / AppFooter.vue / tabs/SettingsPanel.vue / ProfilePanel.vue 用 inject('siteConfig') 直接读取站点配置。"),
    ("动态组件\n(component :is)",
     "<component :is=\"comp\"> 按状态动态渲染不同组件，配合 <KeepAlive> 缓存切换状态。",
     "UserCenter.vue：const currentTab = shallowRef(ProfilePanel) + <KeepAlive><component :is=\"currentTab\" /></KeepAlive> 切换个人信息/地址/设置面板；\nDefaultLayout.vue：<component :is=\"Component\" /> 渲染路由视图。"),
]

table = doc.add_table(rows=1, cols=3)
table.style = "Light Grid Accent 1"
hdr = table.rows[0].cells
for c, t in zip(hdr, ["考核知识点", "说明", "对应实现（文件/组件 + 功能）"]):
    c.paragraphs[0].add_run(t).bold = True
for kp, desc, impl in rows:
    cells = table.add_row().cells
    cells[0].text = kp
    cells[1].text = desc
    cells[2].text = impl
for row in table.rows:
    for cell in row.cells:
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.name = "Microsoft YaHei"
                r._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
                r.font.size = Pt(10)

# ============ 二、模块实现说明 ============
h("二、模块实现说明", 1)

h("模块一：UserCenter.vue — 动态组件 component:is 切换业务面板", 2)
bullet("个人中心用 Tab 在「个人信息 / 收货地址 / 设置」三块业务间切换，主体区用 <component :is=\"currentTab\"> 动态渲染对应面板，并用 <KeepAlive> 缓存切换状态（如地址编辑未保存不丢失）。", bold_prefix="功能简介：")
bullet("动态组件 component:is、shallowRef 持有组件引用、KeepAlive 缓存、v-for 渲染 Tab。", bold_prefix="对应知识点：")
bullet("（此处插入 UserCenter 个人中心页截图）", bold_prefix="效果截图：")

h("模块二：App.vue + AppHeader/AppFooter/SettingsPanel — provide/inject 跨级传值", 2)
bullet("站点配置（站点名、Logo、主题色）在 App.vue 顶层 provide，页头/页脚/设置面板等深层组件直接 inject 使用，主题切换即时生效。", bold_prefix="功能简介：")
bullet("provide/inject 跨级通信、响应式对象跨层共享。", bold_prefix="对应知识点：")
bullet("（此处插入页头 / 设置面板截图）", bold_prefix="效果截图：")

h("模块三：SearchBar.vue — props 校验 + 自定义事件子传父", 2)
bullet("搜索框为子组件，通过 props 接收 modelValue/placeholder 并做类型校验，用户输入后 emit('search', keyword) 把关键词回传父组件 AppHeader 触发路由跳转；同时支持 v-model 双向同步。", bold_prefix="功能简介：")
bullet("props 父传子（类型校验）、defineEmits 自定义事件子传父、组件 v-model。", bold_prefix="对应知识点：")
bullet("（此处插入搜索框截图）", bold_prefix="效果截图：")

h("模块四：ProductDetail / Cart / OrderList — 响应式数据 + 绑定 + 列表/条件渲染 + EP 表格", 2)
bullet("商品详情用 ref/computed 管理当前商品、当前图、数量；购物车/订单用 v-for 与 el-table 渲染列表，用 v-if 按订单状态显示不同操作按钮（去付款/发货/确认收货）。", bold_prefix="功能简介：")
bullet("ref/computed 响应式数据、v-bind 动态 class/style、v-for/v-if 列表与条件渲染、Element Plus el-table。", bold_prefix="对应知识点：")
bullet("（此处插入商品详情 / 购物车 / 订单列表截图）", bold_prefix="效果截图：")

# ============ 三、人员信息 ============
h("三、人员信息", 1)
bullet("吴泓潇", bold_prefix="姓名：")
bullet("商品详情页 ProductDetail.vue（响应式数据、绑定、列表/条件渲染、评价表单 v-model）；\n购物车页 Cart.vue（v-for 列表、事件对象 $event、数量加减）；\n订单列表/详情 OrderList.vue / OrderDetail.vue（el-table、v-if 状态控制）；\n个人中心 UserCenter.vue（动态组件 component:is + KeepAlive）；\n页头/页脚 AppHeader.vue / AppFooter.vue（provide/inject 跨级读取站点配置）；\n搜索框 SearchBar.vue（props 校验 + 自定义事件子传父 + v-model）；\n通用组件 StarRating.vue / CartBadge.vue（props 参数校验）；\n登录页 Login.vue、地址/资料编辑表单（v-model 双向绑定）。", bold_prefix="负责的页面/功能模块：")
bullet("响应式数据、内容/属性绑定、事件绑定、双向绑定、条件 & 列表渲染、父子组件传参、provide/inject、动态组件。", bold_prefix="负责或主要贡献的知识点：")

# ============ 四、问题与解决方案 ============
h("四、开发过程中遇到的问题及解决方案", 1)

para("问题一：深层组件需要共享站点配置，层层 props 传递繁琐且易出错。", bold=True)
bullet("AppHeader、AppFooter、SettingsPanel 等多层组件都要用 siteConfig，若用 props 需逐层透传，耦合高、维护难。", bold_prefix="原因分析：")
bullet("在 App.vue 用 provide('siteConfig', siteConfig) 与 provide('updateSiteConfig', ...) 注入；后代组件用 inject('siteConfig') 直接获取，去掉逐层传参，主题切换即时响应。", bold_prefix="解决方案：")

para("问题二：个人中心切换 Tab 后，已编辑但未保存的表单内容丢失。", bold=True)
bullet("原先用 v-if 条件切换会销毁并重建组件，组件内部状态随之清空。", bold_prefix="原因分析：")
bullet("改用 <component :is=\"currentTab\"> 动态组件，并包裹 <KeepAlive> 缓存切换前的组件实例；currentTab 用 shallowRef 持有组件引用，切换时保留各面板编辑状态。", bold_prefix="解决方案：")

para("问题三：搜索框（子组件）输入的关键词需在页头（父组件）触发页面跳转，值难以回传。", bold=True)
bullet("表单输入状态封装在 SearchBar 内部，父组件拿不到实时值。", bold_prefix="原因分析：")
bullet("SearchBar 用 defineProps 声明 modelValue（类型校验）并 defineEmits(['update:modelValue','search'])；输入时 emit('search', keyword)，父组件 @search=\"onSearch\" 接收并 router.push；对外还支持 v-model 双向绑定。", bold_prefix="解决方案：")

para("问题四（补充）：评价/订单等接口要求登录态，未登录访问会被拦截重定向到登录页。", bold=True)
bullet("请求拦截器对 401 直接跳转登录页，未登录时拉取购物车也会误触发跳转。", bold_prefix="原因分析：")
bullet("在 AppHeader 用 watch 守卫，仅登录后才拉取购物车数据；未登录时不发请求，避免误触 401 重定向。", bold_prefix="解决方案：")

doc.save(OUT)
print("OK ->", OUT)
